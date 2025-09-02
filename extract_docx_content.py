#!/usr/bin/env python3
"""
Script to extract content from Word documents for template analysis
"""
from docx import Document
import json

def extract_document_content(filepath):
    """Extract text content and structure from a Word document"""
    try:
        doc = Document(filepath)
        content = {
            'title': filepath.split('/')[-1],
            'paragraphs': [],
            'tables': [],
            'headers': []
        }
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                style = para.style.name if para.style else 'Normal'
                content['paragraphs'].append({
                    'text': para.text.strip(),
                    'style': style,
                    'is_header': 'Heading' in style
                })
                
                # Track headers separately
                if 'Heading' in style or len(para.text.strip()) < 100 and para.text.strip().endswith(':'):
                    content['headers'].append(para.text.strip())
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            if table_data:
                content['tables'].append({
                    'index': table_idx,
                    'data': table_data
                })
        
        return content
        
    except Exception as e:
        return {'error': str(e), 'title': filepath}

def main():
    documents = [
        "/Users/duncanmiller/Documents/HadadaHealth/Reports/Discharge Summary.docx",
        "/Users/duncanmiller/Documents/HadadaHealth/Reports/Outpatient Planning Record & Report (OPR) 05052025.-2.docx"
    ]
    
    for doc_path in documents:
        print(f"\n{'='*80}")
        print(f"ANALYZING: {doc_path.split('/')[-1]}")
        print(f"{'='*80}")
        
        content = extract_document_content(doc_path)
        
        if 'error' in content:
            print(f"ERROR: {content['error']}")
            continue
            
        print(f"\nHEADERS/SECTIONS FOUND:")
        for header in content['headers']:
            print(f"  • {header}")
            
        print(f"\nSTRUCTURE OVERVIEW:")
        print(f"  • Total paragraphs: {len(content['paragraphs'])}")
        print(f"  • Total tables: {len(content['tables'])}")
        print(f"  • Headers/sections: {len(content['headers'])}")
        
        if content['tables']:
            print(f"\nTABLE STRUCTURES:")
            for i, table in enumerate(content['tables']):
                rows = len(table['data'])
                cols = len(table['data'][0]) if table['data'] else 0
                print(f"  • Table {i+1}: {rows} rows × {cols} columns")
                if table['data']:
                    print(f"    First row: {table['data'][0]}")
        
        print(f"\nKEY CONTENT SECTIONS:")
        for para in content['paragraphs'][:10]:  # Show first 10 paragraphs
            if para['is_header']:
                print(f"  [HEADER] {para['text']}")
            elif len(para['text']) < 200:
                print(f"  {para['text'][:100]}...")
        
        # Save detailed content to file
        output_file = f"template_analysis_{doc_path.split('/')[-1].replace(' ', '_').replace('.docx', '.json')}"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(content, f, indent=2, ensure_ascii=False)
        print(f"\nDetailed analysis saved to: {output_file}")

if __name__ == "__main__":
    main()